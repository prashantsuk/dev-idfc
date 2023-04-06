import os, sys

sys.path.append('../')
import os.path
import shutil
import traceback
import base64
import docx
import docx2txt
from docx.shared import Inches
from PIL import Image
from subprocess import call
from pypdf import PdfReader


class dms_compress():#
    def __init__(self, opt, logger):
        self.opt = opt
        self.logger = logger

    def dtype_format(self, data_dir):
        '''this function will return the file extention and file name'''
        try:
            file_name = data_dir.split('/')[-self.opt.one_val]
            extention_name = file_name.split('.')[-self.opt.one_val]
            return file_name, extention_name
        except Exception as err:
            self.logger.info("unable to find dataformat:{}".format(err))
    def len_of_pdf(self,data_dir):
        reader = PdfReader(data_dir)
        return len(reader.pages)
    def memory_size_eval(self, file):
        size = os.path.getsize(file) / 1000000  # size is in mb
        return size

    def uncompress(self, data_path, destin_dir, file_name):
        destin_dir = destin_dir + file_name
        shutil.move(data_path, destin_dir)
        self.logger.info("unable to compress : {}".format(destin_dir))
        return True,destin_dir, self.memory_size_eval(destin_dir), 1

    def img_compressor(self, img, destin_dir, file_name):
        '''this function will compress all type of image data. refer process_type_checker dict in self.opts.yaml'''
        destin_dir = destin_dir + file_name
        try:
            img = Image.open(img)
            x, y = img.size[self.opt.zero_val], img.size[self.opt.one_val]
            width = int(x * self.opt.scale_percentage / self.opt.cent_val)
            height = int(y * self.opt.scale_percentage / self.opt.cent_val)
            img = img.resize((width, height), Image.ANTIALIAS)
            img.save(destin_dir, optimize=True, quality=self.opt.scale_percentage)
            return True,destin_dir, self.memory_size_eval(destin_dir),1
        except:
            try:  # some image formate need RGB conversion. (format exception handling )
                img = img.convert('RGB')
                img.save(destin_dir, optimize=True, quality=self.opt.scale_percentage)
                return True, destin_dir, self.memory_size_eval(destin_dir),1
            except Exception as err:
                self.logger.info("error unable to compress: {}".format(img))
                self.logger.info("error unable to process img_compressor: {}".format(err))
                return False, destin_dir,0,0

    def get_ghostscript_path(self): # will check the scriptwriter in OS.
        gs_names = ['gs', 'gswin32', 'gswin64']
        for name in gs_names:
            if shutil.which(name):
                return shutil.which(name)
        raise FileNotFoundError(f'No GhostScript executable was found on path ({"/".join(gs_names)})')

    def pdf_compressor(self, data_path, destin_dir, file_name):
        '''this function will compress pdf. refer process_type_checker dict in opts.yaml'''
        pageCount = self.len_of_pdf(data_path)
        try:
            destin_dir = destin_dir + file_name
            quality = self.opt.pdf_quality
            # gs = "/home/azureuser/.local/bin:/home/azureuser/bin:/usr/local/bin:/usr/bin:/usr/local/sbin:/usr/sbin:/usr/lib/jvm/java-11-openjdk-11.0.18.0.10-1.el8_6.x86_64/bin" # To Do: need to put in opt
            gs = self.get_ghostscript_path()
            call([gs, '-sDEVICE=pdfwrite', '-dCompatibilityLevel=1.4',
                  '-dPDFSETTINGS={}'.format(quality[self.opt.zero_val]),
                  '-dNOPAUSE', '-dQUIET', '-dBATCH',
                  '-sOutputFile={}'.format(destin_dir), data_path])
            return True,destin_dir, self.memory_size_eval(destin_dir),pageCount
        except Exception as err:
            self.logger.info("error unable to process pdf_compressor :{}".format(str(err)))

    def image_reader_RGB_cnv(self, path):
        Image.open(path).convert('RGB').save(path, optimize=True, quality=80)

    def document_type_compressor(self, data_path, destin_dir, file_name):
        '''this function will compress word documents. refer process_type_checker dict in opts.yaml'''
        try:
            res_doc = docx.Document()
            image_dirPath = self.opt.image_dirPath  # efs /opt/ path of the local directory\
            destin_dir = destin_dir + file_name
            docx2txt.process(data_path, image_dirPath)
            in_memory_doc = docx.Document(data_path)
            cnt = 0
            par = in_memory_doc.paragraphs
            dp = os.listdir(image_dirPath)
            for i in range(len(par)):
                res_doc.add_paragraph(par[i].text)
                if self.opt.graphicData in par[i]._p.xml:
                    cnt += 1
                    image_name_ext = [i for i in dp if str(cnt) in i]
                    if len(image_name_ext) > 0:  # avo
                        image_path = image_dirPath + image_name_ext[0]
                        self.image_reader_RGB_cnv(image_path)
                        res_doc.add_picture(image_path, width=Inches(5.0))
            res_doc.save(destin_dir)  # compressed data file path
            return True, destin_dir, self.memory_size_eval(destin_dir),1
        except Exception as err:
            self.logger.info("error unable to process document_type_compressor:{}".format(str(err)))
            return False,destin_dir, 0,0

    def dir_gen(self, temp_dir):
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
            self.logger.info("created a local storage : {}".format(temp_dir))

    def base64_to_original_format(self, base64_stream, base64_extention, destin_dir, base64_type, base64_name):
        if base64_extention in base64_type["universal"]: # will convert the base64 data to respective data type
            with open(destin_dir + base64_name + base64_extention, "wb") as f:
                f.write(base64.b64decode(base64_stream))
                f.close()
        if base64_extention in base64_type["text"]:
            file_content = base64.b64decode(base64_stream)
            with open(destin_dir + base64_name + base64_extention, "w+") as f:
                f.write(file_content.decode("utf-8"))
                f.close()
        else:
            self.logger.error("error: unable to decode base64 for :{}".format(base64_extention))
        return destin_dir + base64_name + base64_extention

    def main(self, base64_stream, base64_extention, base64_name):  # main function
        try:
            self.dir_gen(self.opt.destin_dir)  # creating a destin dir in server
            self.dir_gen(self.opt.source_dir) # creating a source dir in server
            self.dir_gen(self.opt.image_dirPath) # creating a image dir in server
            incoming_data = self.base64_to_original_format(base64_stream, base64_extention, self.opt.source_dir,
                                                           self.opt.base64_decrypt, base64_name)
            originalSize = self.memory_size_eval(incoming_data)
            compressedSize, pageCount, compression_status = 0,1,False
            try:
                file_name, extension_name = self.dtype_format(incoming_data)  # getting the file details
                method_call = self.opt.process_type_checker[extension_name]
                dir_compress_data = self.opt.destin_dir  # temperorary local directory::::
                try:
                    if method_call == 'img_compressor':
                        compression_status, destin_dir, compressedSize,pageCount = self.img_compressor(incoming_data, dir_compress_data, file_name)
                    elif method_call == 'pdf_compressor':
                        compression_status,destin_dir, compressedSize,pageCount = self.pdf_compressor(incoming_data, dir_compress_data, file_name)
                    elif method_call == 'document_type_compressor':
                        compression_status, destin_dir,compressedSize,pageCount = self.document_type_compressor(incoming_data, dir_compress_data,file_name)
                    elif method_call == 'uncompress':
                        compression_status,destin_dir, compressedSize,pageCount = self.uncompress(incoming_data, dir_compress_data, file_name)
                    else:
                        self.logger.info("unknown data format detected :{}".format(extension_name))
                    shutil.rmtree(self.opt.source_dir)
                    shutil.rmtree(self.opt.image_dirPath)
                    return originalSize, compression_status,destin_dir, compressedSize, pageCount  # the final compressed memory
                except Exception:
                    traceback.print_exc(file=sys.stdout)
                    self.logger.error("unable to process compression internal function:{}".format(str(traceback.print_exc(file=sys.stdout))))
                    compression_status, destin_dir, compressedSize, pageCount = self.uncompress(incoming_data, dir_compress_data, file_name)
                    return originalSize, compression_status, destin_dir , compressedSize, pageCount
            except Exception:
                traceback.print_exc(file=sys.stdout)
                self.logger.info("error in main_compressor")
                self.logger.error("Unknown file format:{}".format(str(traceback.print_exc(file=sys.stdout))))
                return 0, False, 0
        except Exception as err:
            self.logger.info("error processing base64 conversion: {}" + str(err))
